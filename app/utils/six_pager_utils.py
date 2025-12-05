import streamlit as st
import os
from docx import Document
from io import BytesIO
import base64
import tempfile
from utils.gemini_utils import answer_online
from utils.gpt_utils import answer_online_gpt
from utils.qdrant_utils import get_qdrant_context

collections = ['BCG_energy_transition', 'Towards_sustainable_development_McKinsey', 'McKinsey_report2', 'srccl_ch7_ipcc', 'siemens_sustainability_report'  ]
global_single_collection = "Chanakya" #"strategic_intelligence_test" 
outer_world_reports = [
    "BCGreport_energy_transition",
    "SRCCL_Chapter_7_ipcc",
    "toward-a-sustainable-inclusive-growing-future-the-role-of-business_McKinsey",
    "56297_373_mckinsey_esg_report-2022_aw6_v10_final"
]

def get_firstpage(topic):
    def get_firstpage_from_source(collection_name):
        prompts = [
            f"Introduction to {topic} and its significance"
            # f"Key  {topic} important in understanding them?"
            # f"introduction to {topic} and its relevance?"
            # f"{topic} and its implications in an introductory paragraph?"
            # f"What are the key aspects and implications of {topic}"
            f"How does {topic} contribute to addressing current sustainability goals "
            f"the main objectives or goals associated with {topic}"
            f"primary drivers or motivations behind the development of {topic}"
        ]  # Define prompts specific to the collection here

        context = ""
        images = []
        for prompt in prompts:
            context_from_qdrant, prompt_images = get_qdrant_context(collection_name, prompt)
            
            # print("Response status code:", context_from_qdrant.status_code)
            # print("Response text:", context_from_qdrant.json())

            if context_from_qdrant is not None :
                context += f"{context_from_qdrant}. "
                if len(prompt_images)>0:
                    images.extend(prompt_images[:2]) #Concatenating only first 2 images if they come

        return context, images

    total_context = ""
    images = []
    for collection_name in outer_world_reports:
        context, collection_images = get_firstpage_from_source(collection_name)
        if(len(collection_images)>0):
            images.extend(collection_images)
        total_context += f"{collection_name}: {context}"
    # total_context, images = get_firstpage_from_source(global_single_collection)

    gpt_prompt = f"I want to create a six-pager document. A six-pager is a succinct document spanning six pages, offering a concise overview of a topic or issue. It provides essential insights, analysis, and recommendations in a structured format. Despite its brevity, it aims to convey comprehensive information effectively."
    gpt_prompt = f"The first page of the six pager provides an overview of the topic and its significance, offering a brief introduction to the subject matter. It sets the stage for the reader by highlighting the importance of the topic and its relevance to broader discussions or initiatives."
    gpt_prompt += f"I want to identify the overview and introduction in the {topic} market to fill the introduction part of the six pager document. "
    # gpt_prompt += f"Dump around the {topic} is - {total_context} Generate the introduction for the {topic} from the {total_context} provided."

    final_firstpage = answer_online(gpt_prompt, total_context)
    return final_firstpage, images

def get_secondpage(topic):
    def get_secondpage_from_source(collection_name):
        prompts = [
            f"key milestones and developments in the {topic}",
            # f"What are the different types or categories within the realm of {topic}, and what are their defining characteristics and distinctions?",
            f"What practical guidelines or best practices should be followed in {topic}?",
            # f"How diverse is {topic}, and what are the various types, variants, or subcategories that exist within it?",
            f"case studies or real-life examples of {topic}?",
            f"Start-ups and technologies in the {topic} market"
            f"What are the main variations or subcategories of {topic}",
            f"R&D investments in the {topic}"
            # f"significant trends or patterns in the historical progression of {topic}",
            # f"Role of technological innovation in the development and diversification of {topic}?"
        ]  # Define prompts specific to the collection here

        context = ""
        images = []
        for prompt in prompts:
            context_from_qdrant, prompt_images = get_qdrant_context(collection_name, prompt)
            
            # print("Response status code:", context_from_qdrant.status_code)
            # print("Response text:", context_from_qdrant.json())

            if context_from_qdrant is not None :
                context += f"{context_from_qdrant}. "
                if len(prompt_images)>0:
                    images.extend(prompt_images[:2]) #Concatenating only first 2 images if they come

        return context, images

    total_context = ""
    images = []
    for collection_name in outer_world_reports:
        context, collection_images = get_secondpage_from_source(collection_name)
        if(len(collection_images)>0):
            images.extend(collection_images)
        total_context += f"{collection_name}: {context}"
    # total_context, images = get_secondpage_from_source(global_single_collection)

    gpt_prompt = f"I want to create a six-pager document. A six-pager is a succinct document spanning six pages, offering a concise overview of a topic or issue. It provides essential insights, analysis, and recommendations in a structured format. Despite its brevity, it aims to convey comprehensive information effectively."
    gpt_prompt = f"The second page of the six-pager delves into the historical background of the topic, highlighting key milestones, developments, and evolutionary trends over time. It also discusses the various types or categories within the realm of the topic, outlining their characteristics and distinctions."
    gpt_prompt = f"I want to identify the background and narrative of the {topic} to fill the background part of the six pager document. "
    # gpt_prompt += f"Dump around the {topic} is - {total_context} Generate the background for the {topic} from the {total_context} provided."

    final_secondpage = answer_online(gpt_prompt, total_context)
    return final_secondpage, images

def get_thirdpage(topic):
    def get_thirdpage_from_source(collection_name):
        prompts = [
            f"What are the main obstacles to {topic} development?",
            f"How does {topic} development interact with various factors?",
            f"What new challenges are stakeholders facing in {topic}?",
            f"How do stakeholders address challenges in {topic}?"
            # f"What trends exacerbate challenges in {topic}?",
            # f"How do external factors affect {topic} challenges?",
            f"What role do innovations play in {topic} challenges?",
            f"How do regulations shape {topic} challenges?",
            f"What socio-economic factors influence {topic} challenges?",
            f"Which strategies address {topic} challenges effectively?"
        ]  # Define prompts specific to the collection here

        context = ""
        images = []
        for prompt in prompts:
            context_from_qdrant, prompt_images = get_qdrant_context(collection_name, prompt)
            
            # print("Response status code:", context_from_qdrant.status_code)
            # print("Response text:", context_from_qdrant.json())

            if context_from_qdrant is not None :
                context += f"{context_from_qdrant}. "
                if len(prompt_images)>0:
                    images.extend(prompt_images[:2]) #Concatenating only first 2 images if they come

        return context, images

    total_context = ""
    images = []
    for collection_name in outer_world_reports:
        context, collection_images = get_thirdpage_from_source(collection_name)
        if(len(collection_images)>0):
            images.extend(collection_images)
        total_context += f"{collection_name}: {context}"
    # total_context, images = get_thirdpage_from_source(global_single_collection)

    gpt_prompt = f"I want to create a six-pager document. A six-pager is a succinct document spanning six pages, offering a concise overview of a topic or issue. It provides essential insights, analysis, and recommendations in a structured format. Despite its brevity, it aims to convey comprehensive information effectively."
    gpt_prompt = f"The third page of the six-pager explores the major challenges hindering the development of the topic, considering both internal and external factors. It examines the complexity of topic development by analyzing the interplay of various factors such as technological, socio-economic, and regulatory issues. Additionally, it discusses the evolving landscape of challenges faced by stakeholders, including new threats, opportunities, and uncertainties."
    gpt_prompt = f"I want to identify the problems and challenges of the {topic} to fill the problems part of the six pager document. "
    # gpt_prompt += f"Dump around the {topic} is - {total_context} Generate the probelems and challenges for the {topic} from the {total_context} provided."

    final_thirdpage = answer_online(gpt_prompt, total_context)
    return final_thirdpage, images

def get_fourthpage(topic):
    def get_fourthpage_from_source(collection_name):
        prompts = [
            f"What are existing solutions for {topic} challenges?",
            f"How do emerging technologies aid {topic} challenges?",
            f"What's the role of policy in {topic} development?",
            f"some examples of successful {topic} projects?",
            # f"How do innovations improve {topic} efficiency?",
            # f"What are the components of supportive {topic} policies?",
            f"How do collaborations address {topic} challenges?",
            f"What is the potential of emerging tech in {topic}?",
            # f"Can you give some examples of sustainable {topic} practices?",
            f"What can be the strategies for mitigating {topic} obstacles?"
        ]  # Define prompts specific to the collection here

        context = ""
        images = []
        for prompt in prompts:
            context_from_qdrant, prompt_images = get_qdrant_context(collection_name, prompt)
            
            # print("Response status code:", context_from_qdrant.status_code)
            # print("Response text:", context_from_qdrant.json())

            if context_from_qdrant is not None :
                context += f"{context_from_qdrant}. "
                if len(prompt_images)>0:
                    images.extend(prompt_images[:2]) #Concatenating only first 2 images if they come

        return context, images

    total_context = ""
    images = []
    for collection_name in outer_world_reports:
        context, collection_images = get_fourthpage_from_source(collection_name)
        if(len(collection_images)>0):
            images.extend(collection_images)
        total_context += f"{collection_name}: {context}"
    # total_context, images = get_fourthpage_from_source(global_single_collection)

    gpt_prompt = f"I want to create a six-pager document. A six-pager is a succinct document spanning six pages, offering a concise overview of a topic or issue. It provides essential insights, analysis, and recommendations in a structured format. Despite its brevity, it aims to convey comprehensive information effectively."
    gpt_prompt = f"The fourth page of the six-pager discuss the existing solutions and innovative approaches aimed at mitigating challenges and overcoming obstacles in the field of the topic. It explores the potential of emerging technologies in addressing specific challenges and enhancing efficiency in the sector."
    gpt_prompt = f"I want to identify the solutions to the challenges of the {topic} to fill the solutions part of the six pager document. "
    # gpt_prompt += f"Dump around the {topic} is - {total_context} Generate the solutions for the {topic} from the {total_context} provided."

    final_fourthpage = answer_online(gpt_prompt, total_context)
    return final_fourthpage, images

def get_fifthpage(topic):
    def get_fifthpage_from_source(collection_name):
        prompts = [
            f"What are the key factors influencing the adoption of sustainable practices in {topic} development?",
            f"How can technological innovation contribute to overcoming challenges and achieving sustainability goals in {topic}?",
            f"What role do public-private partnerships play in driving progress and innovation in {topic} development?",
            f"emerging trends or disruptive technologies shaping the future of {topic}?",
            f"How can cross-sector collaboration accelerate progress towards sustainable {topic} development?",
            # f"What measures are being taken to ensure the resilience and adaptability of {topic} initiatives in the face of global challenges?",
            # f"Are there any regulatory barriers hindering the advancement of sustainable practices in {topic} development?",
            f"What opportunities exist for leveraging artificial intelligence in optimizing {topic}",
            f"How can education and awareness campaigns foster a culture of sustainability within {topic} industries?",
            # f"Can you highlight examples of successful {topic} projects and their impact on communities and the environment?"
        ]  # Define prompts specific to the collection here

        context = ""
        images = []
        for prompt in prompts:
            context_from_qdrant, prompt_images = get_qdrant_context(collection_name, prompt)
            
            # print("Response status code:", context_from_qdrant.status_code)
            # print("Response text:", context_from_qdrant.json())

            if context_from_qdrant is not None :
                context += f"{context_from_qdrant}. "
                if len(prompt_images)>0:
                    images.extend(prompt_images[:2]) #Concatenating only first 2 images if they come

        return context, images

    total_context = ""
    images = []
    for collection_name in outer_world_reports:
        context, collection_images = get_fifthpage_from_source(collection_name)
        if(len(collection_images)>0):
            images.extend(collection_images)
        total_context += f"{collection_name}: {context}"
    # total_context, images = get_fifthpage_from_source(global_single_collection)

    gpt_prompt = f"I want to create a six-pager document. A six-pager is a succinct document spanning six pages, offering a concise overview of a topic or issue. It provides essential insights, analysis, and recommendations in a structured format. Despite its brevity, it aims to convey comprehensive information effectively."
    gpt_prompt = f"The fifth page of the six-pager focuses on policy recommendations and strategic interventions aimed at advancing the development of the topic. It takes into account current challenges, emerging trends, and long-term sustainability objectives. Additionally, it identifies priority areas for technological advancement and investment to address critical challenges and unlock new opportunities. "
    gpt_prompt = f"I want to identify the recommendations to the challenges of the {topic} to fill the recommendation part of the six pager document. "
    # gpt_prompt += f"Dump around the {topic} is - {total_context} Generate the recommendations for the {topic} from the {total_context} provided."

    final_fifthpage = answer_online(gpt_prompt, total_context)
    return final_fifthpage, images

def get_sixthpage(topic):
    def get_sixthpage_from_source(collection_name):
        prompts = [
            f"What are the key steps needed to advance {topic} development in the coming years?",
            # f"How can scenario planning enhance {topic} development strategies?",
            f"What risk assessment strategies are crucial for managing {topic} challenges?",
            # f"How to adaptively manage {topic} initiatives in changing environments?",
            f"What are the next steps to catalyze {topic} progress?",
            f"What initiatives can drive impactful change in {topic} development?",
            # f"How to prioritize actions for advancing {topic} goals?",
            f"What strategies can ensure successful implementation of {topic} plans?",
            f"What innovative approaches can accelerate {topic} development?",
            f"How to align {topic} strategies with broader sustainability goals?",
            f"What resources are needed to support {topic} initiatives?"
        ]  # Define prompts specific to the collection here

        context = ""
        images = []
        for prompt in prompts:
            context_from_qdrant, prompt_images = get_qdrant_context(collection_name, prompt)
            
            # print("Response status code:", context_from_qdrant.status_code)
            # print("Response text:", context_from_qdrant.json())

            if context_from_qdrant is not None :
                context += f"{context_from_qdrant}. "
                if len(prompt_images)>0:
                    images.extend(prompt_images[:2]) #Concatenating only first 2 images if they come

        return context, images

    total_context = ""
    images = []
    for collection_name in outer_world_reports:
        context, collection_images = get_sixthpage_from_source(collection_name)
        if(len(collection_images)>0):
            images.extend(collection_images)
        total_context += f"{collection_name}: {context}"
    # total_context, images = get_sixthpage_from_source(global_single_collection)

    gpt_prompt = f"I want to create a six-pager document. A six-pager is a succinct document spanning six pages, offering a concise overview of a topic or issue. It provides essential insights, analysis, and recommendations in a structured format. Despite its brevity, it aims to convey comprehensive information effectively."
    gpt_prompt = f"The sixth page of the six-pager outlines a strategic action plan detailing key steps and initiatives needed to advance the development of the topic in the coming years. It discusses strategies for scenario planning, risk assessment, and adaptive management related to the topic. "
    gpt_prompt = f"I want to identify the plan of action to the challenges of the {topic} to fill the plan of action part of the six pager document. "
    # gpt_prompt += f"Dump around the {topic} is - {total_context} Generate the plan of action for the {topic} from the {total_context} provided."

    final_sixthpage = answer_online(gpt_prompt, total_context)
    return final_sixthpage, images

def generate_six_pager(topic):
   
    first_page_content, first_page_images = get_firstpage(topic)
    second_page_content, second_page_images = get_secondpage(topic)
    third_page_content, third_page_images = get_thirdpage(topic)
    fourth_page_content, fourth_page_images = get_fourthpage(topic)
    fifth_page_content, fifth_page_images = get_fifthpage(topic)
    sixth_page_content, sixth_page_images = get_sixthpage(topic)
    return first_page_content, second_page_content, third_page_content, fourth_page_content, fifth_page_content, sixth_page_content, first_page_images, second_page_images, third_page_images, fourth_page_images, fifth_page_images, sixth_page_images

    # Create a new Word document
    doc = Document()

    # Add content to the document
    doc.add_heading('First Page', level=1)
    doc.add_paragraph(first_page_content)
    add_images_to_doc(first_page_images, doc)

    doc.add_page_break()

    doc.add_heading('Second Page', level=1)
    doc.add_paragraph(second_page_content)
    add_images_to_doc(second_page_images, doc)

    doc.add_page_break()

    doc.add_heading('Third Page', level=1)
    doc.add_paragraph(third_page_content)
    add_images_to_doc(third_page_images, doc)

    doc.add_page_break()

    doc.add_heading('Fourth Page', level=1)
    doc.add_paragraph(fourth_page_content)
    add_images_to_doc(fourth_page_images, doc)

    doc.add_page_break()

    doc.add_heading('Fifth Page', level=1)
    doc.add_paragraph(fifth_page_content)
    add_images_to_doc(fifth_page_images, doc)

    doc.add_page_break()

    doc.add_heading('Sixth Page', level=1)
    doc.add_paragraph(sixth_page_content)
    add_images_to_doc(sixth_page_images, doc)

    # Save the document
    doc.save(f'{topic}_six_pager.docx')
    st.success(f'Six-pager document generated successfully: {topic}_six_pager.docx')

def add_images_to_doc(images:BytesIO, doc):
    for img_bytes in images:
        # Create a temporary file to save the image
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_file.write(img_bytes.getvalue())
        temp_file.close()
        
        # Add the image to the document
        doc.add_picture(temp_file.name)
        
        # Delete the temporary file
        os.unlink(temp_file.name)

def get_binary_file_downloader_html(bin_file, file_label='File', file_name='file.docx'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">{file_label}</a>'
    return href
